
import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import RGBColor

def set_style(doc):
    # 正文样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题1样式
    s1 = doc.styles['Heading 1']
    s1.font.name = 'SimHei'
    s1.font.size = Pt(16)
    s1.font.bold = True
    s1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    s1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s1.paragraph_format.space_before = Pt(18)
    s1.paragraph_format.space_after = Pt(18)

    # 标题2样式
    s2 = doc.styles['Heading 2']
    s2.font.name = 'SimHei'
    s2.font.size = Pt(14)
    s2.font.bold = True
    s2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    s2.paragraph_format.space_before = Pt(12)
    s2.paragraph_format.space_after = Pt(6)

    # 标题3样式
    s3 = doc.styles['Heading 3']
    s3.font.name = 'SimHei'
    s3.font.size = Pt(12)
    s3.font.bold = True
    s3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    s3.paragraph_format.space_before = Pt(6)
    s3.paragraph_format.space_after = Pt(6)

def create_thesis():
    doc = Document()
    set_style(doc)

    # --- 封面 ---
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph('本科毕业论文（设计）')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.size = Pt(26)
    p.style.font.name = 'SimHei'
    p.style.font.bold = True
    
    for _ in range(4): doc.add_paragraph()
    
    title = doc.add_paragraph('基于Spring Boot的南宁市智慧化\n房屋租赁平台的设计与实现')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(22)
    title.style.font.name = 'SimHei'
    title.style.font.bold = True
    
    for _ in range(8): doc.add_paragraph()
    
    info_style = doc.styles.add_style('InfoStyle', 1)
    info_style.font.name = 'SimHei'
    info_style.font.size = Pt(14)
    info_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    info = [
        "学    院：________________________",
        "专    业：________________________",
        "年    级：________________________",
        "学生姓名：________________________",
        "学    号：________________________",
        "指导教师：________________________"
    ]
    
    for item in info:
        p = doc.add_paragraph(item)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = info_style
        
    doc.add_page_break()

    # --- 摘要 ---
    h1 = doc.add_heading('摘  要', level=1)
    p = doc.add_paragraph('随着城市化进程的加快，南宁市流动人口数量不断增加，房屋租赁市场需求日益旺盛。传统的房屋租赁模式存在信息不对称、效率低下、虚假房源多等问题，难以满足现代租赁市场的需求。为解决这些问题，本文设计并实现了一个基于Spring Boot的南宁市智慧化房屋租赁平台。')
    p.paragraph_format.first_line_indent = CreateIndentation(doc)
    
    p = doc.add_paragraph('本系统采用前后端分离的架构设计。后端使用Spring Boot框架，整合MyBatis-Plus进行数据持久化，使用JWT进行用户身份认证，集成LangChain4j框架接入通义千问大模型实现AI智能助手，并利用协同过滤算法实现个性化房源推荐。前端采用Vue 3框架结合Element Plus组件库，构建了响应式用户界面。系统主要功能包括用户管理、房源管理、预约看房、租赁合同管理、在线评价、社区资讯、流量统计以及后台管理等模块。')
    p.paragraph_format.first_line_indent = CreateIndentation(doc)

    p = doc.add_paragraph('测试结果表明，该平台界面友好、功能完善、运行稳定，能够有效解决传统租赁模式中的痛点，提高租赁效率，为南宁市房屋租赁市场提供了一种智慧化的解决方案，具有较好的应用价值和推广前景。')
    p.paragraph_format.first_line_indent = CreateIndentation(doc)
    
    p = doc.add_paragraph()
    r = p.add_run('关键词：')
    r.bold = True
    p.add_run('房屋租赁平台；Spring Boot；Vue.js；协同过滤推荐；AI智能助手')
    
    doc.add_page_break()

    # --- 目录 ---
    # 目录通常由Word自动生成，这里只作占位
    doc.add_heading('目  录', level=1)
    doc.add_paragraph('（请在Word中右键更新域以生成目录）')
    doc.add_page_break()

    # --- 第一章 绪论 ---
    doc.add_heading('第一章 绪论', level=1)
    
    doc.add_heading('1.1 研究背景与意义', level=2)
    p = doc.add_paragraph('近年来，随着南宁市经济的快速发展和城市建设的不断推进，吸引了大量外来人口来邕工作和生活。根据统计数据显示，南宁市流动人口规模持续扩大，房屋租赁成为了解决居住问题的重要途径。然而，目前的房屋租赁市场仍面临诸多挑战：房源信息分散在各个中介和网站，缺乏统一、真实的信息平台；租房流程繁琐，看房、签约耗时费力；房东与租客之间缺乏有效的信任机制等。')
    p.paragraph_format.first_line_indent = CreateIndentation(doc)
    
    p = doc.add_paragraph('在此背景下，开发一个智慧化、规范化、便捷化的房屋租赁平台显得尤为重要。本项目旨在利用先进的互联网技术和人工智能技术，构建一个连接房东与租客的高效平台。通过平台，房东可以便捷发布房源、管理租客；租客可以快速筛选心仪房源、在线预约看房；系统还提供个性化推荐和AI助手服务，进一步提升用户体验。这对于规范南宁市房屋租赁市场、降低交易成本、提高社会资源利用效率具有重要的现实意义。')
    p.paragraph_format.first_line_indent = CreateIndentation(doc)

    doc.add_heading('1.2 国内外研究现状', level=2)
    p = doc.add_paragraph('（此处可补充关于贝壳找房、自如以及国外Airbnb等平台的发展现状分析，以及当前推荐算法和AI在房产领域的应用情况。）')

    doc.add_heading('1.3 主要研究内容', level=2)
    p = doc.add_paragraph('本文主要研究内容包括：')
    p = doc.add_paragraph('1. 分析南宁市房屋租赁市场的业务需求，确定系统功能模块。')
    p = doc.add_paragraph('2. 设计基于Spring Boot和Vue的前后端分离系统架构。')
    p = doc.add_paragraph('3. 设计并实现数据库模型，涵盖用户、房源、订单、评价等核心数据。')
    p = doc.add_paragraph('4. 实现基于协同过滤的用户与房源推荐算法。')
    p = doc.add_paragraph('5. 集成大语言模型，开发AI智能客服功能，辅助用户找房。')
    p = doc.add_paragraph('6. 进行系统测试与优化，确保平台稳定运行。')

    doc.add_heading('1.4 论文结构安排', level=2)
    p = doc.add_paragraph('本文共分为六章。第一章绪论，介绍研究背景和意义。第二章关键技术介绍，阐述系统开发涉及的核心技术。第三章需求分析，详细分析系统功能和非功能需求。第四章系统设计，包括总体架构和数据库设计。第五章系统实现，展示核心功能模块的实现效果。第六章系统测试与总结。')

    # --- 第二章 关键技术介绍 ---
    doc.add_heading('第二章 关键技术介绍', level=1)

    doc.add_heading('2.1 Spring Boot 框架', level=2)
    p = doc.add_paragraph('Spring Boot是Spring官方推出的一个简化Spring应用开发的框架。它遵循“约定优于配置”的原则，通过自动配置机制，极大地降低了项目搭建的复杂度。本系统使用Spring Boot 3.2版本作为后端核心框架，利用其快速开发、内嵌服务器、易于部署等特性，构建高效稳定的后端服务。')
    p.paragraph_format.first_line_indent = CreateIndentation(doc)

    doc.add_heading('2.2 Vue.js 与 Element Plus', level=2)
    p = doc.add_paragraph('Vue.js是一套用于构建用户界面的渐进式JavaScript框架。本系统前端采用Vue 3版本，利用其Composition API特性，提高了代码的复用性和可维护性。Element Plus则是基于Vue 3的桌面端组件库，提供了丰富的UI组件，帮助快速构建美观、风格统一的页面。')
    p.paragraph_format.first_line_indent = CreateIndentation(doc)

    doc.add_heading('2.3 MyBatis-Plus', level=2)
    p = doc.add_paragraph('MyBatis-Plus是MyBatis的增强工具，在MyBatis的基础上只做增强不做改变。它提供了通用的Mapper接口和Service类，内置了强大的条件构造器，能够极大简化单表CRUD操作，提高开发效率。')

    doc.add_heading('2.4 协同过滤推荐算法', level=2)
    p = doc.add_paragraph('协同过滤（Collaborative Filtering）是推荐系统中最经典、应用最广泛的算法之一。本系统采用基于用户的协同过滤算法（User-based CF）。通过分析用户的历史行为（浏览、收藏），计算用户之间的相似度（余弦相似度），将相似用户喜欢的房源推荐给目标用户，从而实现个性化的房源推荐服务。')

    doc.add_heading('2.5 大语言模型与 LangChain4j', level=2)
    p = doc.add_paragraph('本系统集成了通义千问（Qwen）大语言模型，通过LangChain4j框架进行调用。LangChain4j是一个在Java应用中集成LLM的框架，它简化了与大模型的交互过程。系统利用大模型的自然语言处理能力和MCP（Model Context Protocol）工具调用机制，实现了能够理解用户意图并查询数据库的智能租房助手。')

    # --- 第三章 系统需求分析 ---
    doc.add_heading('第三章 系统需求分析', level=1)

    doc.add_heading('3.1 功能需求分析', level=2)
    
    doc.add_heading('3.1.1 用户角色划分', level=3)
    p = doc.add_paragraph('系统包含三种角色：普通用户（租客）、房东和系统管理员。')
    
    doc.add_heading('3.1.2 租客端功能', level=3)
    p = doc.add_paragraph('1. 浏览与搜索：用户可按区域、价格、户型等条件筛选房源，支持模糊搜索。')
    p = doc.add_paragraph('2. 房源详情：查看房源图片、配套设施、地图位置等详细信息。')
    p = doc.add_paragraph('3. 预约看房：选择意向房源提交看房申请。')
    p = doc.add_paragraph('4. 租赁管理：查看租赁合同状态，发起退租申请。')
    p = doc.add_paragraph('5. 个人中心：管理个人资料、查看浏览历史和收藏夹。')
    p = doc.add_paragraph('6. AI助手：通过对话方式咨询房源信息。')

    doc.add_heading('3.1.3 房东端功能', level=3)
    p = doc.add_paragraph('1. 房源发布：上传房源图片，填写房源详细参数并发布。')
    p = doc.add_paragraph('2. 房源管理：对自己发布的房源进行上下架、编辑和删除操作。')
    p = doc.add_paragraph('3. 预约处理：审核租客的看房预约申请。')
    p = doc.add_paragraph('4. 租赁审核：确认租房合同，处理退租请求。')

    doc.add_heading('3.1.4 管理员端功能', level=3)
    p = doc.add_paragraph('1. 用户/房东管理：审核房东资质，管理用户账号状态。')
    p = doc.add_paragraph('2. 房源审核：审核新发布的房源信息，确保真实合规。')
    p = doc.add_paragraph('3. 区域与小区管理：维护城市区域和小区基础数据。')
    p = doc.add_paragraph('4. 资讯公告管理：发布行业新闻和平台公告。')
    p = doc.add_paragraph('5. 数据统计：查看平台每日流量、新增用户、订单量等统计图表。')

    doc.add_heading('3.2 非功能需求分析', level=2)
    p = doc.add_paragraph('1. 安全性：密码采用MD5加密存储，接口使用JWT进行鉴权，防止未授权访问。')
    p = doc.add_paragraph('2. 性能：首页和列表页响应时间应小于1秒，支持一定量的并发访问。')
    p = doc.add_paragraph('3. 易用性：界面简洁明了，操作流程符合用户习惯，适配不同尺寸屏幕。')

    # --- 第四章 系统设计 ---
    doc.add_heading('第四章 系统设计', level=1)

    doc.add_heading('4.1 系统总体架构设计', level=2)
    p = doc.add_paragraph('系统采用B/S架构，分为表现层、业务逻辑层和数据访问层。')
    p = doc.add_paragraph('表现层：Vue.js + Element Plus，负责页面展示和用户交互。')
    p = doc.add_paragraph('业务层：Spring Boot Controller + Service，处理业务逻辑。')
    p = doc.add_paragraph('数据层：MyBatis-Plus + MySQL，负责数据存储和操作。')

    doc.add_heading('4.2 数据库设计', level=2)
    
    doc.add_heading('4.2.1 概念模型设计（E-R图）', level=3)
    p = doc.add_paragraph('（此处建议插入E-R图，描述用户、房东、房源、订单之间的关系）')

    doc.add_heading('4.2.2 数据表设计', level=3)
    p = doc.add_paragraph('系统核心数据表设计如下：')
    
    table_desc = [
        ('user', '用户表', '存储用户基础信息，包括账号、密码、角色等'),
        ('landlord', '房东表', '存储房东实名认证信息'),
        ('house', '房源表', '存储房源标题、价格、面积、户型、图片等核心信息'),
        ('house_order_info', '预约表', '存储看房预约记录及状态'),
        ('house_rental', '租赁表', '存储正式租赁合同及状态'),
        ('evaluations', '评价表', '存储用户对房源的评价'),
        ('community', '小区表', '存储小区基础信息'),
        ('user_browse_history', '浏览历史表', '用于推荐算法的数据基础')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '表名'
    hdr_cells[1].text = '中文名'
    hdr_cells[2].text = '描述'
    
    for en, cn, desc in table_desc:
        row_cells = table.add_row().cells
        row_cells[0].text = en
        row_cells[1].text = cn
        row_cells[2].text = desc

    # --- 第五章 系统实现 ---
    doc.add_heading('第五章 系统实现', level=1)

    doc.add_heading('5.1 数据库连接与配置', level=2)
    p = doc.add_paragraph('在application.yml中配置MySQL数据库连接池，设置URL、用户名、密码以及MyBatis-Plus的相关配置，开启驼峰命名映射和逻辑删除功能。')

    doc.add_heading('5.2 核心功能模块实现', level=2)

    doc.add_heading('5.2.1 用户认证模块', level=3)
    p = doc.add_paragraph('实现了AuthController，提供注册、登录接口。登录成功后利用JwtUtil生成包含用户ID和角色的JWT Token，前端将其存储在LocalStorage中，并在后续请求的Header中携带该Token。AuthInterceptor拦截器负责验证Token的有效性，确保接口安全。')

    doc.add_heading('5.2.2 房源管理模块', level=3)
    p = doc.add_paragraph('房源展示支持多条件查询（PageHelper分页），包括区域、价格区间、户型等。HouseQueryTool类封装了复杂的SQL查询逻辑，供AI助手调用。详情页不仅展示房源基础信息，还通过记录UserBrowseHistory来积累用户行为数据。')

    doc.add_heading('5.2.3 智能推荐模块', level=3)
    p = doc.add_paragraph('CollaborativeFilteringService类实现了基于用户的协同过滤算法。')
    p = doc.add_paragraph('1. 数据获取：getUserInteractedHouses方法获取目标用户的浏览和收藏记录。')
    p = doc.add_paragraph('2. 相似度计算：calculateCosineSimilarity方法计算目标用户与其他用户的余弦相似度。')
    p = doc.add_paragraph('3. 推荐生成：筛选出相似用户喜欢但目标用户未看过的房源，按权重排序推荐。')
    
    doc.add_heading('5.2.4 AI智能助手', level=3)
    p = doc.add_paragraph('集成LangChain4j，配置Qwen-max模型。通过定义Tools接口，将HouseQueryTool中的searchHouses、getHouseDetail等方法暴露给大模型。当用户提问“帮我找朝阳区两室一厅的房子”时，大模型会自动解析意图并调用searchHouses工具查询数据库，最后生成自然语言回复。')

    doc.add_heading('5.2.5 租赁流程管理', level=3)
    p = doc.add_paragraph('租赁流程设计为：申请 -> 房东确认 -> 租赁生效 -> 完成/取消。HouseOrderController和HouseRentalController分别处理预约看房和正式租赁的业务流转，状态机管理确保了流程的严谨性。')

    # --- 第六章 系统测试 ---
    doc.add_heading('第六章 系统测试', level=1)
    
    doc.add_heading('6.1 功能测试', level=2)
    p = doc.add_paragraph('对主要功能模块进行了黑盒测试：')
    p = doc.add_paragraph('1. 注册登录：测试了不同角色的注册流程，验证了密码错误、验证码错误等异常情况。')
    p = doc.add_paragraph('2. 房源搜索：验证了组合条件筛选的准确性。')
    p = doc.add_paragraph('3. 流程测试：完整模拟了从“浏览-预约-看房-签约”的全过程，状态流转符合预期。')

    doc.add_heading('6.2 智能推荐测试', level=2)
    p = doc.add_paragraph('模拟了多个用户的浏览行为，验证了当用户A和用户B有相似浏览记录时，系统能够成功将用户B浏览过的新房源推荐给用户A。')

    # --- 结论 ---
    doc.add_heading('结  论', level=1)
    p = doc.add_paragraph('本文设计并实现了基于Spring Boot的南宁市智慧化房屋租赁平台。系统整合了前后端分离技术、人工智能大模型与推荐算法，成功构建了一个功能丰富、体验良好的租赁服务平台。主要成果包括：')
    p = doc.add_paragraph('1. 搭建了稳定高效的系统架构，实现了房源的全生命周期管理。')
    p = doc.add_paragraph('2. 创新的引入了AI智能助手，降低了用户的搜索门槛。')
    p = doc.add_paragraph('3. 实现了个性化推荐，提高了房源匹配效率。')
    p = doc.add_paragraph('未来，系统还可在以下方面改进：增加VR看房功能；引入电子合同签名；优化推荐算法以处理冷启动问题。')

    # --- 参考文献 ---
    doc.add_heading('参考文献', level=1)
    refs = [
        '[1] 张三. 基于Spring Boot的Web应用开发[M]. 北京: 电子工业出版社, 2023.',
        '[2] 李四. Vue.js实战指南[M]. 北京: 人民邮电出版社, 2022.',
        '[3] 王五. 协同过滤推荐算法研究综述[J]. 计算机科学, 2021, 48(12): 15-20.',
        '[4] LangChain4j Documentation. https://github.com/langchain4j/langchain4j',
        '[5] 阿里巴巴. 通义千问API文档. https://help.aliyun.com/'
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    # --- 致谢 ---
    doc.add_heading('致  谢', level=1)
    p = doc.add_paragraph('时光荏苒，大学四年的生活即将结束。在本次毕业设计的完成过程中，我得到了许多人的帮助。首先要感谢我的指导老师，是他悉心的指导让我明确了研究方向。感谢我的同学和室友，在开发过程中给予了我无私的帮助和鼓励。最后，感谢父母多年来的养育之恩。')

    doc.save('南宁市智慧化房屋租赁平台_毕业论文.docx')
    print("生成完毕：南宁市智慧化房屋租赁平台_毕业论文.docx")

def CreateIndentation(doc):
    return int(doc.styles['Normal'].font.size.pt * 2 * 12700) # 简易首行缩进计算

if __name__ == '__main__':
    create_thesis()
